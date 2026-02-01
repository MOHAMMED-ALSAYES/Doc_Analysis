"""
خدمة استخراج بيانات الطلاب من الوثائق (Excel, Word, PDF, صور)
"""
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

try:
    import pandas as pd
    import openpyxl
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False


class StudentDataExtractor:
    """استخراج بيانات الطلاب من الوثائق"""
    
    def __init__(self):
        self.student_patterns = [
            r'رقم\s*الطالب[:\s]*(\d+)',
            r'الطالب[:\s]*(\d+)',
            r'student\s*number[:\s]*(\d+)',
            r'id[:\s]*(\d+)',
        ]
        
        self.name_patterns = [
            r'الاسم[:\s]*([^\n\r]+)',
            r'اسم\s*الطالب[:\s]*([^\n\r]+)',
            r'name[:\s]*([^\n\r]+)',
        ]
        
        self.score_patterns = [
            r'الدرجة[:\s]*(\d+(?:\.\d+)?)',
            r'النقاط[:\s]*(\d+(?:\.\d+)?)',
            r'score[:\s]*(\d+(?:\.\d+)?)',
            r'grade[:\s]*(\d+(?:\.\d+)?)',
            r'(\d+(?:\.\d+)?)\s*/\s*(\d+(?:\.\d+)?)',  # 85/100
        ]
    
    def extract_from_excel(self, file_path: Path) -> List[Dict[str, Any]]:
        """استخراج بيانات الطلاب من ملف Excel"""
        if not PANDAS_AVAILABLE:
            return []
        
        students_data = []
        
        try:
            # قراءة ملف Excel
            df = pd.read_excel(file_path, engine='openpyxl')
            
            # البحث عن الأعمدة المحتملة
            student_number_col = None
            name_col = None
            score_col = None
            subject_col = None
            
            # البحث في أسماء الأعمدة
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['رقم', 'number', 'id', 'student']):
                    if 'student' in col_lower or 'رقم' in col_lower:
                        student_number_col = col
                if any(keyword in col_lower for keyword in ['اسم', 'name', 'الاسم']):
                    name_col = col
                if any(keyword in col_lower for keyword in ['درجة', 'score', 'grade', 'mark', 'نقاط']):
                    score_col = col
                if any(keyword in col_lower for keyword in ['مادة', 'subject', 'course']):
                    subject_col = col
            
            # إذا لم نجد أعمدة محددة، نستخدم الأعمدة الأولى
            if student_number_col is None and len(df.columns) > 0:
                student_number_col = df.columns[0]
            if name_col is None and len(df.columns) > 1:
                name_col = df.columns[1]
            if score_col is None and len(df.columns) > 2:
                score_col = df.columns[2]
            
            # استخراج البيانات من كل صف
            for idx, row in df.iterrows():
                try:
                    student_number = None
                    name = None
                    score = None
                    subject = None
                    
                    # استخراج رقم الطالب
                    if student_number_col:
                        val = row.get(student_number_col)
                        if pd.notna(val):
                            student_number = str(int(float(val))) if isinstance(val, (int, float)) else str(val).strip()
                    
                    # استخراج الاسم
                    if name_col:
                        val = row.get(name_col)
                        if pd.notna(val):
                            name = str(val).strip()
                    
                    # استخراج الدرجة
                    if score_col:
                        val = row.get(score_col)
                        if pd.notna(val):
                            if isinstance(val, (int, float)):
                                score = float(val)
                            else:
                                # محاولة استخراج رقم من النص
                                match = re.search(r'(\d+(?:\.\d+)?)', str(val))
                                if match:
                                    score = float(match.group(1))
                    
                    # استخراج المادة
                    if subject_col:
                        val = row.get(subject_col)
                        if pd.notna(val):
                            subject = str(val).strip()
                    
                    # إضافة فقط إذا كان لدينا رقم طالب واسم
                    if student_number and name:
                        students_data.append({
                            'student_number': student_number,
                            'full_name': name,
                            'score': score,
                            'subject': subject,
                            'row_index': idx + 1,
                        })
                except Exception as e:
                    print(f"[WARN] خطأ في استخراج بيانات الصف {idx + 1}: {e}")
                    continue
            
        except Exception as e:
            print(f"[ERROR] خطأ في قراءة ملف Excel: {e}")
        
        return students_data
    
    def extract_from_text(self, text: str) -> List[Dict[str, Any]]:
        """استخراج بيانات الطلاب من النص المستخرج (OCR)"""
        students_data = []
        
        if not text:
            return students_data
        
        # تقسيم النص إلى أسطر وتنظيفها
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # محاولة استخراج من تنسيق الجدول بـ pipes (|)
        # مثال: "الرقم | الاسم | الدرجة" أو "1 | محمد سايس | 100"
        table_format_detected = False
        header_line_index = -1
        
        # البحث عن سطر العنوان (header) الذي يحتوي على pipes
        # نبحث عن كلمات مفتاحية متعددة بما في ذلك "مجموع الدرجات"، "النسبة"، "التقدير"
        # وأسماء المواد المختلفة (حاسوب، محاسبة، إنجليزي، دورات)
        header_keywords = [
            'الرقم', 'الاسم', 'الدرجة', 'number', 'name', 'grade', 
            'مجموع', 'النسبة', 'التقدير', 'ملاحظات',
            'total', 'percentage', 'notes', 'score', 'marks',
            # مواد الحاسوب
            'windows', 'word', 'powerpoint', 'excel', 'انترنت', 'اکسس', 'access', 'internet',
            'icdl', 'office', 'outlook', 'publisher',
            # مواد المحاسبة
            'محاسبة', 'accounting', 'مالية', 'financial', 'اقتصاد', 'economy',
            'قيد', 'entry', 'ميزانية', 'budget', 'تكاليف', 'cost',
            # مواد اللغة الإنجليزية
            'انجليزي', 'english', 'reading', 'writing', 'grammar', 'vocabulary',
            'قواعد', 'مفردات', 'قراءة', 'كتابة', 'listening', 'speaking',
            'حضور', 'مشاركة', 'واجبات', 'تحدث', 'استماع', 'نصفي', 'نهائي',
            'attendance', 'participation', 'assignments', 'midterm', 'final',
            # كلمات عامة للجداول
            'مادة', 'subject', 'course', 'دورة', 'module', 'وحدة'
        ]
        
        for idx, line in enumerate(lines):
            if '|' in line:
                line_lower = line.lower()
                # إذا كان السطر يحتوي على pipes وكلمة مفتاحية، فهو header محتمل
                if any(keyword.lower() in line_lower for keyword in header_keywords) or \
                   any(keyword in line for keyword in header_keywords):
                    # التحقق من أنه ليس صف بيانات (لا يبدأ برقم عادة)
                    parts = line.split('|')
                    first_part = parts[0].strip() if parts else ""
                    # header عادة لا يبدأ برقم صحيح
                    if not first_part.isdigit() and len(first_part) < 10:
                        table_format_detected = True
                        header_line_index = idx
                        print(f"[DEBUG] تم اكتشاف header في السطر {idx}: {line[:100]}")
                        break
        
        # إذا لم نجد header واضح، نبحث عن أي سطر يحتوي على pipes وبيانات تبدو كجدول
        if not table_format_detected:
            for idx, line in enumerate(lines):
                if '|' in line:
                    parts = [p.strip() for p in line.split('|')]
                    # إذا كان السطر يحتوي على 3+ أعمدة وأول عمود رقم، قد يكون جدول
                    if len(parts) >= 3:
                        first_part = parts[0].strip()
                        if first_part.isdigit() and int(first_part) > 0 and int(first_part) < 1000:
                            # التحقق من وجود اسم في العمود التالي
                            if len(parts) > 1 and parts[1].strip() and not parts[1].strip().isdigit():
                                table_format_detected = True
                                header_line_index = idx - 1 if idx > 0 else -1  # قد يكون header قبل صف واحد
                                break
        
        # إذا وجدنا تنسيق جدول، استخراج البيانات من الأسطر التالية
        if table_format_detected:
            # استخراج أسماء المواد من header
            subjects_map = {}  # {column_index: subject_name}
            if header_line_index >= 0:
                header_line = lines[header_line_index]
                header_parts = [p.strip() for p in header_line.split('|')]
                
                # كلمات نريد تجاهلها (ليست مواد)
                ignore_words = [
                    # عناوين عامة
                    'الرقم', 'الاسم', 'الدرجة', 'number', 'name', 'grade', 'م',
                    'مجموع', 'النسبة', 'التقدير', 'ملاحظات', 'notes',
                    'total', 'percentage', 'appreciation', 'appreciate',
                    'score', 'scores', 'marks', 'mark', 'grade',
                    # كلمات إضافية
                    'seq', 'serial', 'no', 'nr', 'id', 'student', 'طالب',
                    'result', 'results', 'exam', 'test', 'quiz'
                ]
                
                for i, part in enumerate(header_parts):
                    part_clean = part.lower().strip()
                    # إذا كان الجزء ليس فارغاً وليس في قائمة الكلمات المستبعدة
                    if part_clean and part_clean not in [w.lower() for w in ignore_words]:
                        # إذا كان يحتوي على نص (ليس رقم فقط)، قد يكون اسم مادة
                        if not part_clean.replace('.', '').isdigit() and len(part_clean) > 1:
                            subjects_map[i] = part.strip()
                            print(f"[DEBUG] مادة في العمود {i}: {part.strip()}")
            
            # البدء من السطر التالي للـ header
            start_idx = header_line_index + 1 if header_line_index >= 0 else 0
            
            for idx in range(start_idx, len(lines)):
                line = lines[idx]
                if '|' not in line:
                    continue
                
                # تقسيم السطر بـ pipes
                parts = [p.strip() for p in line.split('|')]
                
                if len(parts) >= 2:
                    # محاولة تحديد الرقم والاسم والدرجات
                    student_number = None
                    full_name = None
                    subject_grades = []  # قائمة الدرجات للمواد
                    total_score = None
                    percentage = None
                    max_score = 100
                    
                    # البحث عن رقم الطالب في الأعمدة الأولى
                    for i in range(min(3, len(parts))):
                        part = parts[i].strip()
                        if not part:
                            continue
                        # رقم الطالب: رقم صحيح صغير (عادة 1-100)
                        if part.isdigit():
                            num_val = int(part)
                            if 1 <= num_val <= 999 and not student_number:
                                student_number = part
                                break
                    
                    # البحث عن الاسم في الأعمدة التالية
                    # عادة يكون الاسم في العمود الثاني بعد رقم الطالب
                    name_start_idx = 1 if student_number and student_number == parts[0].strip() else 0
                    for i in range(name_start_idx, min(name_start_idx + 5, len(parts))):
                        part = parts[i].strip()
                        if not part or part == student_number:
                            continue
                        
                        # تخطي الأعمدة الفارغة المتتالية
                        if not part and i < len(parts) - 1 and not parts[i+1].strip():
                            continue
                        
                        # إذا كان نص وليس رقم فقط، قد يكون الاسم
                        part_no_commas = part.replace('.', '').replace('-', '').replace('،', '').replace(',', '').replace('|', '').strip()
                        if part_no_commas and not part_no_commas.isdigit():
                            # تنظيف الاسم
                            name_clean = re.sub(r'[^\w\s\u0600-\u06FF]', ' ', part).strip()
                            # تجاهل الكلمات الشائعة والكلمات القصيرة جداً
                            if len(name_clean) > 2 and name_clean not in ['الرقم', 'الاسم', 'الدرجة', 'Number', 'Name', 'Grade', 
                                                                          'مجموع', 'النسبة', 'التقدير', 'ملاحظات', 'م', 'notes']:
                                # إذا كان الاسم أطول أو لم يكن لدينا اسم بعد، استخدمه
                                if not full_name or (len(name_clean) > len(full_name) and len(name_clean) > 5):
                                    full_name = name_clean
                                    # إذا وجدنا اسم جيد، نتوقف عن البحث (عادة الاسم في عمود واحد)
                                    if len(full_name) > 10:
                                        break
                    
                    # استخراج درجات المواد الفردية
                    # نستخرج جميع المواد من الـ header، حتى إذا كانت فارغة
                    for col_idx, subject_name in subjects_map.items():
                        # تحديد الدرجة الكاملة من اسم المادة (مثل "(10)" أو "(20)")
                        max_score_for_subject = 100  # افتراض 100 كقيمة افتراضية
                        score_match = re.search(r'\((\d+)\)', subject_name)
                        if score_match:
                            max_score_for_subject = int(score_match.group(1))
                        
                        if col_idx < len(parts):
                            part = parts[col_idx].strip()
                            
                            # إذا كانت الخلية فارغة، نسجل أن الطالب لم يختبر المادة
                            if not part:
                                subject_grades.append({
                                    'subject': subject_name,
                                    'score': None,
                                    'max_score': max_score_for_subject,
                                    'percentage': None,
                                    'notes': 'لم يختبر'
                                })
                                print(f"[DEBUG] مادة {subject_name}: لم يختبر (فارغة)")
                                continue
                            
                            # التحقق من القيم الخاصة (مثل "غ" التي تعني غائب)
                            part_lower = part.lower()
                            if part_lower in ['غ', 'gh', 'absent', 'غائب', 'غير', 'لا', 'no', 'n/a', 'na', 'ضعيف']:
                                # نسجل أن الطالب لم يختبر المادة
                                subject_grades.append({
                                    'subject': subject_name,
                                    'score': None,
                                    'max_score': max_score_for_subject,
                                    'percentage': None,
                                    'notes': 'لم يختبر' if part_lower != 'ضعيف' else 'ضعيف'
                                })
                                print(f"[DEBUG] مادة {subject_name}: لم يختبر (قيمة خاصة: {part})")
                                continue
                            
                            try:
                                # تنظيف الجزء من أي أحرف غير رقمية (لكن نسمح بالنقاط للكسور العشرية)
                                clean_part = re.sub(r'[^\d.]', '', part)
                                if clean_part:
                                    score_val = float(clean_part)
                                    # نطاق الدرجات: 0-200
                                    if 0 <= score_val <= 200:
                                        percentage_for_subject = score_val
                                        
                                        # حساب النسبة المئوية بناءً على الدرجة الكاملة
                                        if max_score_for_subject > 0:
                                            percentage_for_subject = (score_val / max_score_for_subject) * 100
                                        
                                        if score_val > 100 and max_score_for_subject == 100:
                                            # إذا كانت الدرجة أكثر من 100 ولم نجد الدرجة الكاملة في الاسم
                                            # قد تكون من 150 أو 200
                                            if score_val <= 150:
                                                max_score_for_subject = 150
                                                percentage_for_subject = (score_val / 150) * 100
                                            else:
                                                max_score_for_subject = 200
                                                percentage_for_subject = (score_val / 200) * 100
                                        
                                        subject_grades.append({
                                            'subject': subject_name,
                                            'score': score_val,
                                            'max_score': max_score_for_subject,
                                            'percentage': percentage_for_subject,
                                        })
                                        print(f"[DEBUG] مادة {subject_name}: {score_val} من {max_score_for_subject} ({percentage_for_subject:.2f}%)")
                                else:
                                    # إذا لم نستطع تحويل إلى رقم، نسجل كـ "لم يختبر"
                                    subject_grades.append({
                                        'subject': subject_name,
                                        'score': None,
                                        'max_score': max_score_for_subject,
                                        'percentage': None,
                                        'notes': 'لم يختبر'
                                    })
                                    print(f"[DEBUG] مادة {subject_name}: لم يختبر (قيمة غير رقمية: {part})")
                            except Exception as e:
                                # في حالة خطأ، نسجل كـ "لم يختبر"
                                subject_grades.append({
                                    'subject': subject_name,
                                    'score': None,
                                    'max_score': max_score_for_subject,
                                    'percentage': None,
                                    'notes': 'لم يختبر'
                                })
                                print(f"[DEBUG] خطأ في استخراج درجة المادة {subject_name}: {e} - سجل كـ 'لم يختبر'")
                        else:
                            # إذا كان العمود غير موجود في الصف، نسجل كـ "لم يختبر"
                            subject_grades.append({
                                'subject': subject_name,
                                'score': None,
                                'max_score': max_score_for_subject,
                                'percentage': None,
                                'notes': 'لم يختبر'
                            })
                            print(f"[DEBUG] مادة {subject_name}: لم يختبر (العمود غير موجود)")
                    
                    # البحث عن مجموع الدرجات والنسبة في الأعمدة الأخيرة
                    scores_found = []
                    for i in range(len(parts)):
                        part = parts[i].strip()
                        if not part:
                            continue
                        
                        # تخطي القيم الخاصة (مثل "غ")
                        part_lower = part.lower()
                        if part_lower in ['غ', 'gh', 'absent', 'غائب', 'ضعيف', 'weak']:
                            continue
                        
                        # محاولة تحويل إلى رقم
                        try:
                            clean_part = re.sub(r'[^\d.]', '', part)
                            if clean_part:
                                score_val = float(clean_part)
                                # نطاق الدرجات: 0-1000
                                if 0 <= score_val <= 1000:
                                    # تجاهل الأرقام الصغيرة جداً التي قد تكون أرقام طلاب (لكن نسمح بـ 0 للمجموع)
                                    if (score_val > 5 and score_val <= 100) or (score_val > 100) or (score_val == 0 and i >= len(parts) - 3):
                                        scores_found.append((i, score_val))
                        except:
                            pass
                    
                    # البحث عن مجموع الدرجات والنسبة في الأعمدة الأخيرة
                    if scores_found:
                        for i, score_val in reversed(scores_found):
                            # التحقق من سياق العمود (مجموع أو نسبة)
                            context_text = ' '.join(parts[max(0, i-1):min(len(parts), i+2)]).lower()
                            
                            # البحث في header أيضاً لمعرفة نوع العمود
                            header_context = ""
                            if header_line_index >= 0 and i < len(header_parts):
                                header_context = header_parts[i].lower() if i < len(header_parts) else ""
                            
                            combined_context = f"{context_text} {header_context}"
                            
                            if i >= len(parts) - 3:  # في الأعمدة الأخيرة
                                if 'مجموع' in combined_context or 'total' in combined_context or 'رقماً' in combined_context or 'كتابة' in combined_context:
                                    if score_val > 0 or not total_score:  # نفضل القيم الموجبة
                                        total_score = score_val
                                elif 'النسبة' in combined_context or '%' in combined_context or 'نسبة' in combined_context:
                                    if score_val > 0 or not percentage:
                                        percentage = score_val
                                elif score_val > 100:
                                    # إذا كانت الدرجة أكثر من 100، قد تكون مجموع
                                    total_score = score_val
                                elif score_val <= 100 and not percentage:
                                    # إذا كانت 100 أو أقل ولم نجد نسبة بعد، قد تكون نسبة
                                    percentage = score_val
                    
                    # إذا كان لدينا رقم طالب واسم، أضف الطالب
                    if student_number and full_name:
                        # إذا كان لدينا درجات للمواد، ننشئ سجلاً لكل مادة
                        if subject_grades:
                            for grade_info in subject_grades:
                                students_data.append({
                                    'student_number': student_number,
                                    'full_name': full_name,
                                    'subject': grade_info['subject'],
                                    'score': grade_info['score'],
                                    'max_score': grade_info['max_score'],
                                    'percentage': grade_info['percentage'],
                                })
                        
                        # إذا كان لدينا مجموع درجات أو نسبة إجمالية، نضيفها كسجل إضافي
                        if total_score is not None or percentage is not None:
                            students_data.append({
                                'student_number': student_number,
                                'full_name': full_name,
                                'subject': 'المجموع الكلي' if total_score else 'النسبة الإجمالية',
                                'score': total_score if total_score else None,
                                'max_score': total_score if total_score else 100,
                                'percentage': percentage if percentage else (total_score / len(subject_grades) if total_score and subject_grades else None),
                            })
                        # إذا لم يكن لدينا درجات للمواد ولكن لدينا درجة عامة
                        elif not subject_grades and scores_found:
                            # نستخدم أول درجة صالحة
                            score_val = scores_found[-1][1]
                            students_data.append({
                                'student_number': student_number,
                                'full_name': full_name,
                                'subject': None,  # بدون مادة محددة
                                'score': score_val,
                                'max_score': score_val if score_val > 100 else 100,
                                'percentage': score_val if score_val <= 100 else None,
                            })
                        
                        print(f"[DEBUG] تم استخراج: رقم={student_number}, اسم={full_name}, مواد={len(subject_grades)}, مجموع={total_score}, نسبة={percentage}")
            
            # إذا استخرجنا بيانات من الجدول، نعيدها
            if students_data:
                print(f"[INFO] تم استخراج {len(students_data)} طالب من تنسيق الجدول")
                return students_data
        
        # إذا لم نجد تنسيق جدول، نستخدم الطريقة القديمة
        current_student = {}
        
        for line in lines:
            # البحث عن رقم الطالب
            for pattern in self.student_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    current_student['student_number'] = match.group(1).strip()
                    break
            
            # البحث عن الاسم
            for pattern in self.name_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    name = match.group(1).strip()
                    # تنظيف الاسم من علامات الترقيم الزائدة
                    name = re.sub(r'[^\w\s\u0600-\u06FF]', '', name)
                    current_student['full_name'] = name
                    break
            
            # البحث عن الدرجة
            for pattern in self.score_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    if len(match.groups()) == 2:
                        # صيغة 85/100
                        score = float(match.group(1))
                        max_score = float(match.group(2))
                        current_student['score'] = score
                        current_student['max_score'] = max_score
                        current_student['percentage'] = (score / max_score) * 100 if max_score > 0 else None
                    else:
                        score = float(match.group(1))
                        current_student['score'] = score
                        current_student['max_score'] = 100
                        current_student['percentage'] = score
                    break
            
            # إذا كان لدينا بيانات كافية، أضف الطالب
            if 'student_number' in current_student and 'full_name' in current_student:
                students_data.append(current_student.copy())
                current_student = {}
        
        return students_data
    
    def extract_from_word(self, file_path: Path) -> str:
        """استخراج النص من ملف Word"""
        if not DOCX_AVAILABLE:
            return ""
        
        try:
            doc = docx.Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            # استخراج النص من الجداول أيضاً
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as e:
            print(f"[WARN] خطأ في قراءة ملف Word: {e}")
            return ""
    
    def extract_from_pdf(self, file_path: Path) -> str:
        """استخراج النص من ملف PDF"""
        text = ""
        
        # محاولة مع PyMuPDF أولاً (أسرع)
        if PDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                text_parts = []
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                text = "\n".join(text_parts)
                if text.strip():
                    return text
            except Exception as e:
                print(f"[WARN] خطأ في قراءة PDF مع PyMuPDF: {e}")
        
        # محاولة مع pdfminer كبديل
        if PDFMINER_AVAILABLE and not text.strip():
            try:
                text = pdfminer_extract(file_path)
            except Exception as e:
                print(f"[WARN] خطأ في قراءة PDF مع pdfminer: {e}")
        
        return text
    
    def detect_has_grades(self, text: str) -> bool:
        """الكشف عن وجود درجات في النص"""
        if not text:
            return False
        
        text_lower = text.lower()
        
        # كلمات مفتاحية للدرجات
        grade_keywords = [
            'درجة', 'نقاط', 'score', 'grade', 'mark', 'total',
            'كشف درجات', 'جدول درجات', 'نتائج',
            'اختبار', 'امتحان', 'test', 'exam'
        ]
        
        # البحث عن كلمات مفتاحية
        has_keywords = any(keyword in text_lower for keyword in grade_keywords)
        
        # البحث عن أنماط درجات (رقم/رقم أو رقم مئوي)
        score_pattern = re.compile(r'\d+(?:\.\d+)?\s*/?\s*\d{1,3}(?:\s*%|درجة|نقطة)?', re.IGNORECASE)
        has_score_pattern = bool(score_pattern.search(text))
        
        # البحث عن أرقام طلاب متعددة
        student_pattern = re.compile(r'(?:رقم|number|id)[\s:]*(\d+)', re.IGNORECASE)
        student_matches = student_pattern.findall(text)
        has_multiple_students = len(set(student_matches)) >= 2
        
        return (has_keywords or has_score_pattern) and has_multiple_students
    
    def extract_from_file(self, file_path: Path, ocr_text: Optional[str] = None) -> List[Dict[str, Any]]:
        """استخراج بيانات الطلاب من ملف (Excel, Word, PDF أو نص)"""
        file_ext = file_path.suffix.lower()
        
        students_data = []
        extracted_text = None
        
        # محاولة استخراج من Excel أولاً
        if file_ext in ['.xlsx', '.xls']:
            students_data = self.extract_from_excel(file_path)
        
        # محاولة استخراج النص من Word
        if file_ext in ['.docx', '.doc']:
            extracted_text = self.extract_from_word(file_path)
            if extracted_text:
                students_data = self.extract_from_text(extracted_text)
        
        # محاولة استخراج النص من PDF
        if file_ext == '.pdf' and not students_data:
            extracted_text = self.extract_from_pdf(file_path)
            if extracted_text:
                students_data = self.extract_from_text(extracted_text)
        
        # إذا لم نحصل على بيانات، جرب النص من OCR
        if not students_data and ocr_text:
            students_data = self.extract_from_text(ocr_text)
        
        return students_data
    
    def check_document_classification(self, text: str, students_data: List[Dict[str, Any]]) -> Optional[str]:
        """تحديد تصنيف الوثيقة بناءً على المحتوى"""
        # إذا احتوت على درجات طلاب، تصنف كـ "كشف درجات"
        if students_data and len(students_data) > 0:
            # التحقق من وجود درجات في البيانات
            has_grades = any(
                s.get('score') is not None or 
                s.get('percentage') is not None 
                for s in students_data
            )
            if has_grades:
                return 'كشف درجات'
        
        # التحقق من وجود كلمات مفتاحية للدرجات في النص
        if self.detect_has_grades(text):
            return 'كشف درجات'
        
        # التحقق من كلمات مفتاحية للاختبار
        if text:
            test_keywords = ['اختبار', 'امتحان', 'test', 'exam', 'quiz', 'سؤال', 'question']
            text_lower = text.lower()
            if any(keyword in text_lower for keyword in test_keywords):
                return 'اختبار'
        
        return None
    
    def calculate_grade(self, percentage: float) -> str:
        """حساب التقدير بناءً على النسبة المئوية"""
        if percentage >= 90:
            return 'A'
        elif percentage >= 80:
            return 'B'
        elif percentage >= 70:
            return 'C'
        elif percentage >= 60:
            return 'D'
        else:
            return 'F'


# إنشاء instance عام
student_extractor = StudentDataExtractor()
 