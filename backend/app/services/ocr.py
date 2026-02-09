"""
خدمة OCR خفيفة للخطة المجانية - تستخدم Tesseract فقط
"""
from pathlib import Path
from typing import Tuple
import io
import re
from datetime import datetime

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document as DocxDocument
import openpyxl


def extract_text_from_pdf(pdf_path: Path, langs: str = "ara+eng") -> Tuple[str, float]:
    """استخراج النص من ملف PDF"""
    doc = fitz.open(str(pdf_path))
    full_text_parts: list[str] = []
    total_pages = len(doc)
    total_confidence = 0.0

    for page_num, page in enumerate(doc):
        # محاولة استخراج النص النصي أولاً (دقة 100%)
        text = page.get_text("text", flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE)
        
        # تحسين استخراج الجداول من PDF
        try:
            tables = page.find_tables()
            if tables:
                for table in tables:
                    table_text = table.extract()
                    if table_text:
                        table_rows = []
                        for row in table_text:
                            if row:
                                row_text = " | ".join(str(cell).strip() if cell else "" for cell in row)
                                if row_text.strip():
                                    table_rows.append(row_text)
                        if table_rows:
                            full_text_parts.append("\n".join(table_rows))
        except Exception:
            pass
        
        if text and text.strip():
            full_text_parts.append(text)
            total_confidence += 100.0
            continue
        
        # إذا لم يوجد نص نصي، استخدم OCR
        try:
            pix = page.get_pixmap(dpi=300, alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            ocr_text, ocr_accuracy = ocr_image_to_text(img, langs=langs)
            if ocr_text:
                full_text_parts.append(ocr_text)
                total_confidence += ocr_accuracy
        except Exception as e:
            print(f"[WARN] OCR error on page {page_num}: {e}")

    combined = "\n".join(full_text_parts).strip()
    
    if total_pages > 0:
        accuracy = total_confidence / total_pages
    else:
        accuracy = 0.0
    
    combined = _post_process_text(combined)
    
    return combined, round(accuracy, 2)


def extract_text_from_word(docx_path: Path) -> Tuple[str, float]:
    """استخراج النص من ملف Word (.docx أو .doc)"""
    file_path = Path(docx_path)
    suffix = file_path.suffix.lower()
    
    if suffix == '.doc':
        return _extract_text_from_old_doc(file_path)
    
    try:
        doc = DocxDocument(str(file_path))
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cell_text = re.sub(r'\s+', ' ', cell_text)
                    row_cells.append(cell_text if cell_text else "")
                
                row_text = " | ".join(row_cells)
                if row_text.strip():
                    table_rows.append(row_text)
            
            if table_rows:
                full_text.extend(table_rows)
        
        combined = "\n".join(full_text).strip()
        
        if not combined:
            return "", 0.0
        
        return combined, 100.0
    except Exception as e:
        print(f"[ERROR] خطأ في استخراج النص من Word: {e}")
        if suffix == '.docx':
            return _extract_text_from_old_doc(file_path)
        return "", 0.0


def _extract_text_from_old_doc(doc_path: Path) -> Tuple[str, float]:
    """استخراج النص من ملف .doc القديم"""
    try:
        import subprocess
        
        try:
            result = subprocess.run(
                ['antiword', str(doc_path)],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore',
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                text = re.sub(r'\t+', ' | ', text)
                text = re.sub(r'\s+', ' ', text)
                return text, 95.0
        except FileNotFoundError:
            pass
        except Exception:
            pass
        
        try:
            result = subprocess.run(
                ['catdoc', str(doc_path)],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore',
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                text = re.sub(r'\t+', ' | ', text)
                text = re.sub(r'\s+', ' ', text)
                return text, 90.0
        except FileNotFoundError:
            pass
        except Exception:
            pass
        
        return "", 0.0
            
    except Exception as e:
        print(f"[ERROR] خطأ في معالجة ملف .doc: {e}")
        return "", 0.0


def extract_text_from_excel(xlsx_path: Path) -> Tuple[str, float]:
    """استخراج النص من ملف Excel"""
    try:
        suffix = xlsx_path.suffix.lower()
        
        if suffix == '.xls':
            try:
                import xlrd
                workbook = xlrd.open_workbook(str(xlsx_path))
                full_text = []
                
                for sheet_idx in range(workbook.nsheets):
                    sheet = workbook.sheet_by_index(sheet_idx)
                    full_text.append(f"--- {sheet.name} ---")
                    
                    for row_idx in range(sheet.nrows):
                        row_values = []
                        for col_idx in range(sheet.ncols):
                            cell = sheet.cell(row_idx, col_idx)
                            cell_value = str(cell.value) if cell.value is not None else ""
                            row_values.append(cell_value.strip())
                        
                        row_text = " | ".join(row_values)
                        if row_text.strip():
                            full_text.append(row_text)
                
                combined = "\n".join(full_text).strip()
                return combined, 100.0
            except ImportError:
                pass
            except Exception:
                pass
        
        wb = openpyxl.load_workbook(str(xlsx_path), data_only=True, read_only=True)
        full_text = []
        
        for sheet in wb.worksheets:
            full_text.append(f"--- {sheet.title} ---")
            
            for row in sheet.iter_rows(values_only=True):
                row_values = []
                for cell in row:
                    if cell is None:
                        row_values.append("")
                    elif isinstance(cell, (int, float)):
                        if isinstance(cell, float) and cell.is_integer():
                            row_values.append(str(int(cell)))
                        else:
                            row_values.append(str(cell))
                    else:
                        row_values.append(str(cell).strip())
                
                row_text = " | ".join(row_values)
                if row_text.strip():
                    full_text.append(row_text)
        
        combined = "\n".join(full_text).strip()
        return combined, 100.0
    except Exception as e:
        print(f"[ERROR] خطأ في استخراج النص من Excel: {e}")
        return "", 0.0


def extract_text_from_image(image_path: Path, langs: str = "ara+eng") -> Tuple[str, float]:
    """استخراج النص من صورة"""
    try:
        img = Image.open(str(image_path))
        return ocr_image_to_text(img, langs=langs)
    except Exception as e:
        print(f"خطأ في استخراج النص من الصورة: {e}")
        return "", 0.0


from PIL import Image, ImageOps, ImageFilter

def ocr_image_to_text(pil_image: Image.Image, langs: str = "ara+eng") -> Tuple[str, float]:
    """تشغيل Tesseract OCR على الصورة مع تحسين الجودة"""
    try:
        # 1. تحويل إلى Scala رمادية (Grayscale)
        if pil_image.mode != 'L':
             gray_img = pil_image.convert('L')
        else:
             gray_img = pil_image
        
        # 2. تحسين التباين تلقائياً (Auto Contrast)
        # هذا يساعد جداً في الصور الملتقطة بالكاميرا حيث الإضاءة غير متساوية
        enhanced_img = ImageOps.autocontrast(gray_img, cutoff=2)
        
        # 3. تطبيق فلتر لزيادة الحدة (Sharpen) لتوضيح الحروف
        sharpened_img = enhanced_img.filter(ImageFilter.SHARPEN)
        
        # 4. تكبير الصورة قليلاً إذا كانت صغيرة (لتحسين دقة Tesseract)
        width, height = sharpened_img.size
        if width < 1500:
             factor = 2
             sharpened_img = sharpened_img.resize((width * factor, height * factor), Image.Resampling.LANCZOS)

        # استخدام Tesseract
        # psm 3: Fully automatic page segmentation, but no OSD. (Default)
        # psm 6: Assume a single uniform block of text.
        # للكاميرا والصور العشوائية، psm 3 أو 1 أفضل.
        custom_config = r'--oem 3 --psm 3' 
        
        text = pytesseract.image_to_string(sharpened_img, lang=langs, config=custom_config)
        
        # حساب دقة تقريبية بناءً على طول النص
        if text.strip():
            processed_text = _post_process_text(text)
            word_count = len(processed_text.split())
            # معادلة بسيطة: كلما زاد عدد الكلمات، زادت الثقة (إلى حد ما)
            accuracy = min(95.0, 40.0 + (word_count / 5)) 
            return processed_text, round(accuracy, 2)
            
    except Exception as e:
        print(f"[ERROR] Tesseract error: {e}")
        # محاولة احتياطية بدون معالجة في حال فشل الفلاتر
        try:
             text = pytesseract.image_to_string(pil_image, lang=langs)
             return text, 50.0
        except:
             pass
    
    return "", 0.0


def _post_process_text(text: str) -> str:
    """معالجة نهائية للنص المستخرج"""
    if not text:
        return ""
    
    # إزالة الأحرف المكررة
    text = re.sub(r'(ال){2,}', 'ال', text)
    text = re.sub(r'([\u0600-\u06FF])\1{2,}', r'\1\1', text)
    
    # إزالة المسافات المتعددة
    text = re.sub(r'[ \t]{2,}', ' ', text)
    
    # إضافة مسافة بعد علامات الترقيم
    text = re.sub(r'([.!?؛،:])([^\s\n])', r'\1 \2', text)
    
    # إزالة المسافات قبل علامات الترقيم
    text = re.sub(r'\s+([.!?؛،:])', r'\1', text)
    
    # تنظيف الأسطر
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    text = '\n'.join(lines)
    
    # إزالة الأسطر الفارغة المتعددة
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def extract_text_smart(file_path: Path, langs: str = "ara+eng") -> Tuple[str, float]:
    """
    استخراج ذكي للنص حسب نوع الملف
    Returns: (text, accuracy_percentage)
    """
    suffix = file_path.suffix.lower()
    
    if suffix in ['.doc', '.docx']:
        text, accuracy = extract_text_from_word(file_path)
        if not text or accuracy < 50:
            print(f"[WARN] استخراج النص من Word فشل، نحاول PDF")
        return text, accuracy
    
    elif suffix == '.pdf':
        return extract_text_from_pdf(file_path, langs)
    
    elif suffix in ['.xlsx', '.xls']:
        return extract_text_from_excel(file_path)
    
    elif suffix in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
        return extract_text_from_image(file_path, langs)
    
    elif suffix == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text, 100.0
        except Exception as e:
            print(f"[ERROR] خطأ في قراءة الملف النصي: {e}")
            return "", 0.0
    
    else:
        print(f"[WARN] نوع ملف غير مدعوم: {suffix}")
        return "", 0.0
